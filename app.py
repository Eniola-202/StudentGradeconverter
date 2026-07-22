import pandas as pd
from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH


def calculate_grade(score):
  "Mathematical computation for grade classification ."
  try:
    score= float(score)
  except ValueError:
    return 'F', 0.0, "_"
  if score >=75:
    return "A", 4, score
  elif score >=70:
    return "AB", 3.5, score
  elif score >=65:
    return "B", 3.25, score
  elif score >=60:
    return "BC", 3.00, score
  elif score >=55:
    return "C", 2.75, score
  elif score >=50:
    return "CD", 2.50, score
  elif score >=45:
    return "D", 2.25, score
  elif score >=40:
    return "E", 2.00, score
  else:
    return "F", 0.0, score


def generaete_transcript(excel_file_path, out_doc_path, credit_unit_row_index=0):
  print(f"Reading data from {excel_file_path}")
  df= pd.read_excel(excel_file_path, skiprows=4)
  # departmental_path= pd.read_excel(excel_file_path, skiprows=1) to get departmental name
  print("These are the exact columns Python sees:", df.columns.tolist())
  doc=Document()

  # grouped_students= df.groupby('NAME ')

  df.columns= df.columns.str.strip()

  start_idx = df.columns.get_loc('COURSE CODE') + 1
  end_idx = df.columns.get_loc('TGP')
  course_cols= df.columns[start_idx:end_idx]
  credit_units= df.iloc[credit_unit_row_index]
  student_data= df.drop(index=credit_unit_row_index)

  for index, row in student_data.iterrows():
    student_name= str(row['NAME'])
    remark_col= str(row['REMARK'])
    if remark_col.lower() == 'fail' or not remark_col.strip():
      continue

    if student_name.lower() =='nan' or not student_name.strip():
      continue
    student_matric_no= str(row['MATRIC NO'])
    if student_matric_no.lower() =='nan' or not student_matric_no.strip():
      continue
    tgp= str(row['TGP'])
    if tgp.lower() =='nan' or not tgp.strip():
      continue
    

    #Heading for doc
    heading= doc.add_heading('STUDENT TRANSCRIPT', level=1)
    heading.alignment= WD_ALIGN_PARAGRAPH.CENTER

    #add student data
    doc.add_paragraph(f"Name: {student_name}")
    doc.add_paragraph(f"Matric Number: {student_matric_no}")
    # doc.add_paragraph(f"Faculty: {faculty}")
    # doc.add_paragraph(f"Department: {department}")
    doc.add_paragraph()

    #create table for courses
    table= doc.add_table(rows=1, cols=6)
    table.style= 'Table Grid'

    #Table Headers
    hdr_cells= table.rows[0].cells
    hdr_cells[0].text= 'S/N'
    hdr_cells[1].text= 'Course Code'
    hdr_cells[2].text= 'Credit Unit'
    hdr_cells[3].text= 'Grade'

    total_credit_units=0
    total_grade_points=0.00

    counter= 1
    credit_units_row = df.iloc[0]

    for course in course_cols:
      score = row[course]
      cu_raw = credit_units[course]

      if pd.isna(score):
        continue

      try: 
        credit_unit = float(cu_raw)
      except (ValueError, TypeError):
        credit_unit =0.0

      grade, grade_weight, remark = calculate_grade(score)

      course_grade_point= credit_unit * grade_weight
      credit_unit= float(credit_units_row[course])
      total_credit_units += credit_unit
      total_grade_points += course_grade_point
      df_indexed = df.set_index('NAME')
      grade_student= df_indexed.loc[student_name, course]

      row_cells = table.add_row().cells
      row_cells[0].text= str(counter)
      row_cells[1].text= str(course)
      row_cells[2].text= str(credit_unit)
      row_cells[3].text= str(grade_student)
      counter += 1

    if total_credit_units > 0:
      gpa= float(tgp)/float(total_credit_units)
    else:
      gpa= 0.0

    doc.add_paragraph() #empty line
    summay_paragraph= doc.add_paragraph()
    summay_paragraph.add_run("Result Summary\n").bold =True
    summay_paragraph.add_run(f"Total Credit Units: {total_credit_units}\n")
    summay_paragraph.add_run(f"Total Grade Points: {tgp}\n")
    summay_paragraph.add_run(f"GPA: {gpa:.2f}\n").bold= True
    summay_paragraph.add_run(f"Remark: {remark_col}\n")
    doc.add_page_break()

  doc.save(out_doc_path)
  print(f"success compile and saved to {out_doc_path}")

if __name__== "__main__":
  INPUT_EXCEL= "/content/first_semester.xlsx"
  OUTPUT_DOC= "/content/Compiled_transcripts.docx"

generaete_transcript(INPUT_EXCEL, OUTPUT_DOC)